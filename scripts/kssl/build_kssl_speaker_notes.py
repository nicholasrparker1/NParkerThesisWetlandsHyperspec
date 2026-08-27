"""Build concise speaker notes matched to the current complete-results deck."""

from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs/reports/KSSL_hydric_project_complete_results_speaker_notes.docx"

NOTES = [
    ("Integrating Laboratory Soil Spectroscopy with Spatial Wetland Evidence", "This project tests whether laboratory MIR spectra and measured soil properties can strengthen interpretation of remotely sensed wetland soils. The current work establishes the laboratory and spatial evidence framework needed before airborne imagery is added."),
    ("Study workflow keeps evidence sources independent", "The workflow moves from the full KSSL archive to a surface-soil regional cohort, MIR processing, and increasingly strict validation. Laboratory spectra are tested against independent mapped evidence rather than being used to define the reference categories."),
    ("The KSSL archive supports a traceable analysis", "The audit confirms a large archive with measured properties, MIR scans, and generally four technical replicates per spectral master. This establishes sufficient data volume while preserving measurement method and replicate provenance."),
    ("Laboratory properties form coherent multivariate gradients", "The correlation and PCA graphics show that the extracted chemistry behaves coherently, but samples form continuous gradients rather than distinct hydric classes. Laboratory chemistry is therefore informative evidence, not a stand-alone hydric label."),
    ("Spatial overlays connect laboratory samples to mapped wetland evidence", "This map connects 404 Montana and North Dakota surface samples to SSURGO hydric-soil context and NWI wetlands. It creates the geographic reference needed to compare laboratory measurements with independent landscape evidence."),
    ("Independent spatial evidence focuses the study on 404 surface samples", "The graphics summarize how often SSURGO and NWI agree or disagree across the regional cohort. Agreement strengthens the weak reference signal, while disagreement identifies cases that require caution or additional evidence."),
    ("The version-0 tool exposes evidence strength and its construction", "The diagnostic shows that most samples have low mapped evidence and makes the SSURGO/NWI score construction transparent. Input completeness is reported separately from scientific certainty, so a complete record is not mistaken for a confirmed hydric determination."),
    ("Laboratory properties overlap across spatial-evidence groups", "These distributions show extensive overlap in carbon, clay, pH, water retention, and related properties among mapped-evidence groups. This supports a multivariate, uncertainty-aware tool rather than fixed thresholds on one property."),
    ("A quality-controlled MIR cohort is analysis-ready", "The mean spectra summarize 398 usable samples after resolving rescans and averaging technical replicates. Group differences are subtle relative to within-group variability, but the spectra provide a stable basis for multivariate modeling."),
    ("MIR variation is continuous—not a simple hydric split", "MIR PCA captures major spectral variation, but the mapped-evidence groups overlap strongly. The spectra contain compositional information, although there is no simple universal hydric spectral cluster."),
    ("MIR strongly predicts laboratory bridge properties", "Project-grouped validation shows strong prediction of carbon, clay, CEC, pH, and water retention from MIR spectra. These interpretable properties are the most defensible bridge between laboratory MIR and future airborne observations."),
    ("Spatial hydric evidence is detectable within project-grouped validation", "MIR has a modest association with the provisional mapped-evidence score when entire KSSL projects are held out. The signal is much weaker than soil-property prediction and should be treated as exploratory."),
    ("Geographic holdouts change the hydric interpretation", "Training in one state and predicting the other shows that several soil-property rankings transfer, but the mapped hydric-evidence association does not. This prevents us from claiming a general MIR-only hydric delineation model."),
    ("What the evidence supports—and the next bridge", "The supported result is that MIR recovers soil-property gradients relevant to wetland interpretation; the unsupported claim is universal hydric classification. Airborne VNIR–SWIR data will be connected through shared soil properties, colocated samples, and spatial context—not by assuming direct wavelength-to-wavelength equivalence with MIR."),
    ("Technical appendix", "The appendix documents secondary diagnostics, coverage, quality control, and modeling safeguards. These slides support questions without interrupting the main results narrative."),
    ("A1. Depth is a first-order remote-sensing constraint", "The boxplots show how key properties vary with layer depth using explicit units. Because airborne imagery observes the exposed surface, surface and subsurface layers must remain separated."),
    ("A2. Candidate evidence is not a complete hydric label", "This figure shows the number of samples supported by taxonomy or related candidate signals before spatial overlays. Indeterminate records are unknown, not automatically non-hydric."),
    ("A3. KSSL sampling is geographically extensive", "The national map shows the breadth and uneven distribution of KSSL sampling. It describes archive coverage rather than wetland extent or a prediction surface."),
    ("A4. Data coverage varies across spatial groups", "This graphic reports property and MIR availability within each evidence group. Unequal coverage and group size are retained explicitly because they affect model reliability."),
    ("A5. Technical replicate QC", "The replicate diagnostic measures disagreement among repeated MIR scans for each sample. Strong replicate consistency supports averaging them into one spectrum per selected sample master."),
    ("A6. MIR PCA loadings", "The loadings identify the spectral regions contributing most to the first two MIR components. They guide interpretation and feature review but are not chemical assignments by themselves."),
    ("A7. Project-grouped observed-versus-predicted results", "Each prediction comes from a fold that excludes the sample's complete KSSL project. This is a stricter and more realistic test than randomly splitting technical replicates or closely related samples."),
]


def main() -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("KSSL Complete Results — Speaker Notes", 0)
    doc.add_paragraph("One- to two-sentence explanations matched to the current 22-slide presentation.")
    for number, (title, note) in enumerate(NOTES, start=1):
        doc.add_heading(f"Slide {number}: {title}", level=1)
        doc.add_paragraph(note)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
